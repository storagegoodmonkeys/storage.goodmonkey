#!/usr/bin/env python3
"""
Create professionally formatted Word document and PDF for client response
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.platypus import Table, TableStyle
from reportlab.lib import colors
from reportlab.platypus.flowables import HRFlowable
import os

def create_word_document():
    """Create professional Word document"""
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Response to Additional Feature Requests', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_heading('Flick MVP Development', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Header info
    header_info = [
        "To: Good Monkeys LLC",
        "From: Saddam - Founder, CodeFlow Studios", 
        "Date: January 10, 2025",
        "Subject: Response to Additional Feature Requests - Flick MVP Development"
    ]
    
    for info in header_info:
        p = doc.add_paragraph(info)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add spacing
    doc.add_paragraph()
    doc.add_paragraph("─" * 80)
    doc.add_paragraph()
    
    # Opening paragraph
    opening = doc.add_paragraph()
    opening.add_run("Dear Good Monkeys LLC Team,\n\n")
    opening.add_run("Thank you for these excellent feature suggestions! Your vision for the Flick app shows great forward-thinking. I'm pleased to address each of your requests and explain how we'll approach them within our MVP scope while setting the foundation for future expansion.\n\n")
    
    # Section 1
    doc.add_heading('1. QR Code Scanning Enhancement ✅ INCLUDED', level=2)
    
    p1 = doc.add_paragraph()
    p1.add_run("Your QR code behavior specification is ")
    p1.add_run("excellent").bold = True
    p1.add_run(" and we're excited to implement this comprehensive flow:\n\n")
    
    p1.add_run("✅ MVP Implementation:\n").bold = True
    features1 = [
        "• Native iOS camera QR scanning with deep linking to the app",
        "• Web-based QR landing pages for non-app users",
        "• Smart routing based on lighter registration status",
        "• \"Lost/Found\" decision interface for registered lighters",
        "• \"Download app\" promotion for unregistered lighters",
        "• In-app QR scanner remains available for registered users"
    ]
    
    for feature in features1:
        doc.add_paragraph(feature)
    
    p1_end = doc.add_paragraph()
    p1_end.add_run("\nThis creates a seamless user experience that encourages app downloads while providing immediate value to anyone who finds a lighter.\n\n")
    
    # Section 2
    doc.add_heading('2. Ownership History Feature ✅ INCLUDED', level=2)
    
    p2 = doc.add_paragraph()
    p2.add_run("The ownership history concept adds tremendous value to the community aspect of Flick:\n\n")
    p2.add_run("✅ MVP Implementation:\n").bold = True
    
    features2 = [
        "• Ownership chain tracking in our database schema",
        "• Privacy-respecting display (only public profiles shown)",
        "• Clean timeline view: \"X → A → B → You\"",
        "• Optional privacy settings for each user",
        "• Integration with existing trade/lost & found workflows"
    ]
    
    for feature in features2:
        doc.add_paragraph(feature)
    
    p2_end = doc.add_paragraph()
    p2_end.add_run("\nThis feature will make each lighter's journey feel more personal and engaging for users.\n\n")
    
    # Section 3
    doc.add_heading('3. Location Update Feature ✅ INCLUDED', level=2)
    
    p3 = doc.add_paragraph()
    p3.add_run("Your proactive location tracking approach is innovative:\n\n")
    p3.add_run("✅ MVP Implementation:\n").bold = True
    
    features3 = [
        "• Smart push notifications asking \"Is your lighter with you?\"",
        "• One-tap responses (Yes/No)",
        "• Automatic location updates when user confirms",
        "• Follow-up \"Is it lost?\" flow for missing lighters",
        "• Non-intrusive frequency (max 1-2 prompts per week)",
        "• Respects user privacy and battery life"
    ]
    
    for feature in features3:
        doc.add_paragraph(feature)
    
    p3_end = doc.add_paragraph()
    p3_end.add_run("\nThis will significantly improve the lost & found success rate while keeping the experience user-friendly.\n\n")
    
    # Section 4
    doc.add_heading('4. Social Feed Architecture ✅ FUTURE-READY', level=2)
    
    p4 = doc.add_paragraph()
    p4.add_run("✅ MVP Foundation:\n").bold = True
    p4.add_run("We'll structure our backend and database to support social features from day one:\n\n")
    
    foundation_items = [
        "• Database Design: Activity log tables for trades, lost/found events, achievements",
        "• API Architecture: Event-driven system ready for feed generation",
        "• User Interaction Tracking: All actions logged for future social features",
        "• Scalable Infrastructure: Built to handle increased social activity"
    ]
    
    for item in foundation_items:
        doc.add_paragraph(item)
    
    p4_future = doc.add_paragraph()
    p4_future.add_run("\nFuture Expansion Ready:\n").bold = True
    p4_future.add_run("When you're ready to add social feeds, we'll have all the data and infrastructure in place for a smooth implementation.\n\n")
    
    # Section 5
    doc.add_heading('5. goodmonkeys.com Landing Website ✅ ENHANCED SCOPE', level=2)
    
    p5 = doc.add_paragraph()
    p5.add_run("Your dual-purpose website concept is perfect for brand building:\n\n")
    p5.add_run("✅ MVP Implementation:\n").bold = True
    
    website_features = [
        "• Company Section: About Good Monkeys LLC, mission, team",
        "• Flick App Section: App introduction, features showcase, download buttons",
        "• Contact Form: Professional contact system with email forwarding",
        "• Responsive Design: Mobile-optimized for all devices",
        "• SEO Optimized: Ready for search engine visibility",
        "• Integrated Experience: Seamless flow between company info and app promotion"
    ]
    
    for feature in website_features:
        doc.add_paragraph(feature)
    
    p5_end = doc.add_paragraph()
    p5_end.add_run("\nThis creates a professional presence that serves both your company brand and Flick app marketing needs.\n\n")
    
    # Divider
    doc.add_paragraph("─" * 80)
    doc.add_paragraph()
    
    # Updated Scope Summary
    doc.add_heading('📋 Updated MVP Scope Summary', level=2)
    
    p_scope = doc.add_paragraph()
    p_scope.add_run("All your requested features are now included in our $13,500 MVP development:\n\n").bold = True
    
    # Core Features
    doc.add_heading('Core Features:', level=3)
    core_features = [
        "✅ iOS App with native QR scanning + web fallback",
        "✅ User profiles with privacy controls",
        "✅ Ownership history tracking",
        "✅ Smart location update system",
        "✅ Lost & found workflow",
        "✅ Trading & gifting system",
        "✅ Basic gamification",
        "✅ Admin dashboard",
        "✅ Enhanced goodmonkeys.com website"
    ]
    
    for feature in core_features:
        doc.add_paragraph(feature)
    
    # Technical Foundation
    doc.add_heading('Technical Foundation:', level=3)
    tech_foundation = [
        "✅ Social-ready database architecture",
        "✅ Scalable API design",
        "✅ Event logging for future features",
        "✅ Privacy-first user controls"
    ]
    
    for item in tech_foundation:
        doc.add_paragraph(item)
    
    # Development Approach
    doc.add_heading('🚀 Development Approach', level=2)
    
    approach_para = doc.add_paragraph()
    approach_para.add_run("MVP Philosophy: ").bold = True
    approach_para.add_run("We'll implement these features in their ")
    approach_para.add_run("core, functional form").bold = True
    approach_para.add_run(" to deliver maximum value while maintaining our 12-week timeline. Each feature will be:\n\n")
    
    approach_items = [
        "• Fully functional for immediate use",
        "• User-friendly with intuitive interfaces", 
        "• Scalable for future enhancement",
        "• Privacy-compliant with user controls"
    ]
    
    for item in approach_items:
        doc.add_paragraph(item)
    
    future_para = doc.add_paragraph()
    future_para.add_run("\nFuture Expansion: ").bold = True
    future_para.add_run("Our architecture ensures that when you're ready to add advanced social features, video content, or complex analytics, the transition will be smooth and cost-effective.\n\n")
    
    # Timeline & Deliverables
    doc.add_heading('📅 Timeline & Deliverables', level=2)
    
    timeline_para = doc.add_paragraph()
    timeline_para.add_run("No changes to our timeline").bold = True
    timeline_para.add_run(" - all features will be delivered within our original 12-week schedule:\n\n")
    
    timeline_items = [
        "• Weeks 1-2: Planning & enhanced scope design",
        "• Weeks 3-8: Core development including all new features",
        "• Weeks 9-10: Testing & optimization",
        "• Weeks 11-12: Deployment & launch"
    ]
    
    for item in timeline_items:
        doc.add_paragraph(item)
    
    # Investment Confirmation
    doc.add_heading('💰 Investment Confirmation', level=2)
    
    investment_para = doc.add_paragraph()
    investment_para.add_run("Total cost remains $13,500").bold = True
    investment_para.add_run(" - we're excited to include these valuable features at no additional cost because they align perfectly with our vision for Flick's success.\n\n")
    
    # Next Steps
    doc.add_heading('🤝 Next Steps', level=2)
    
    next_steps = [
        "1. Confirm approval of this enhanced scope",
        "2. Schedule kickoff call to finalize technical details", 
        "3. Begin development November 1-7, 2025"
    ]
    
    for step in next_steps:
        doc.add_paragraph(step)
    
    # Closing
    closing_para = doc.add_paragraph()
    closing_para.add_run("\nWe're genuinely excited about these enhancements - they'll make Flick a truly innovative and engaging platform that users will love. Your vision for community-driven lighter tracking is going to create something special.\n\n")
    closing_para.add_run("Thank you for trusting CodeFlow Studios with this exciting project!\n\n")
    
    # Signature
    doc.add_paragraph("Best regards,")
    doc.add_paragraph()
    doc.add_paragraph("Saddam")
    doc.add_paragraph("Founder, CodeFlow Studios")
    doc.add_paragraph("Pakistan")
    doc.add_paragraph("📱 WhatsApp: +92-333-4443355")
    doc.add_paragraph("🌐 Website: https://www.codeflowstudios.xyz/")
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph("─" * 80)
    
    footer_para = doc.add_paragraph()
    footer_para.add_run("P.S. - We're already brainstorming some additional creative features we could add in future versions. The foundation we're building will support some amazing expansions!").italic = True
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save Word document
    word_filename = "/Users/tayyab/Desktop/flick/Client_Response_Feature_Requests.docx"
    doc.save(word_filename)
    print(f"✅ Word document created: {word_filename}")
    return word_filename


def create_pdf_document():
    """Create professional PDF document"""
    
    filename = "/Users/tayyab/Desktop/flick/Client_Response_Feature_Requests.pdf"
    doc = SimpleDocTemplate(
        filename,
        pagesize=A4,
        rightMargin=1*inch,
        leftMargin=1*inch,
        topMargin=1*inch,
        bottomMargin=1*inch
    )
    
    story = []
    styles = getSampleStyleSheet()
    
    # Custom colors
    primary_blue = HexColor('#667eea')
    dark_gray = HexColor('#2c3e50')
    text_gray = HexColor('#333333')
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=20,
        alignment=TA_CENTER,
        textColor=primary_blue,
        fontName='Helvetica-Bold',
        leading=30
    )
    
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Heading2'],
        fontSize=16,
        spaceAfter=15,
        alignment=TA_CENTER,
        textColor=dark_gray,
        fontName='Helvetica-Bold',
        leading=20
    )
    
    section_style = ParagraphStyle(
        'CustomSection',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=12,
        alignment=TA_LEFT,
        textColor=primary_blue,
        fontName='Helvetica-Bold',
        leading=18
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=8,
        alignment=TA_JUSTIFY,
        textColor=text_gray,
        fontName='Helvetica',
        leading=14
    )
    
    bullet_style = ParagraphStyle(
        'CustomBullet',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=4,
        alignment=TA_LEFT,
        textColor=text_gray,
        fontName='Helvetica',
        leftIndent=20,
        leading=13
    )
    
    # Title
    story.append(Paragraph("Response to Additional Feature Requests", title_style))
    story.append(Paragraph("Flick MVP Development", subtitle_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Header info
    header_data = [
        ['To:', 'Good Monkeys LLC'],
        ['From:', 'Saddam - Founder, CodeFlow Studios'],
        ['Date:', 'January 10, 2025'],
        ['Subject:', 'Response to Additional Feature Requests - Flick MVP Development']
    ]
    
    header_table = Table(header_data, colWidths=[1*inch, 5*inch])
    header_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP')
    ]))
    
    story.append(header_table)
    story.append(Spacer(1, 0.2*inch))
    story.append(HRFlowable(width="100%", thickness=1, color=dark_gray))
    story.append(Spacer(1, 0.1*inch))
    
    # Opening paragraph
    story.append(Paragraph("Dear Good Monkeys LLC Team,", normal_style))
    story.append(Spacer(1, 0.05*inch))
    story.append(Paragraph(
        "Thank you for these excellent feature suggestions! Your vision for the Flick app shows great forward-thinking. I'm pleased to address each of your requests and explain how we'll approach them within our MVP scope while setting the foundation for future expansion.",
        normal_style
    ))
    story.append(Spacer(1, 0.15*inch))
    
    # Section 1
    story.append(Paragraph("1. QR Code Scanning Enhancement ✅ INCLUDED", section_style))
    story.append(Paragraph(
        "Your QR code behavior specification is <b>excellent</b> and we're excited to implement this comprehensive flow:",
        normal_style
    ))
    story.append(Spacer(1, 0.05*inch))
    
    qr_features = [
        "• Native iOS camera QR scanning with deep linking to the app",
        "• Web-based QR landing pages for non-app users",
        "• Smart routing based on lighter registration status",
        "• \"Lost/Found\" decision interface for registered lighters",
        "• \"Download app\" promotion for unregistered lighters",
        "• In-app QR scanner remains available for registered users"
    ]
    
    for feature in qr_features:
        story.append(Paragraph(feature, bullet_style))
    
    story.append(Paragraph(
        "This creates a seamless user experience that encourages app downloads while providing immediate value to anyone who finds a lighter.",
        normal_style
    ))
    story.append(Spacer(1, 0.15*inch))
    
    # Section 2
    story.append(Paragraph("2. Ownership History Feature ✅ INCLUDED", section_style))
    story.append(Paragraph(
        "The ownership history concept adds tremendous value to the community aspect of Flick:",
        normal_style
    ))
    story.append(Spacer(1, 0.05*inch))
    
    ownership_features = [
        "• Ownership chain tracking in our database schema",
        "• Privacy-respecting display (only public profiles shown)",
        "• Clean timeline view: \"X → A → B → You\"",
        "• Optional privacy settings for each user",
        "• Integration with existing trade/lost & found workflows"
    ]
    
    for feature in ownership_features:
        story.append(Paragraph(feature, bullet_style))
    
    story.append(Paragraph(
        "This feature will make each lighter's journey feel more personal and engaging for users.",
        normal_style
    ))
    story.append(Spacer(1, 0.15*inch))
    
    # Section 3
    story.append(Paragraph("3. Location Update Feature ✅ INCLUDED", section_style))
    story.append(Paragraph(
        "Your proactive location tracking approach is innovative:",
        normal_style
    ))
    story.append(Spacer(1, 0.05*inch))
    
    location_features = [
        "• Smart push notifications asking \"Is your lighter with you?\"",
        "• One-tap responses (Yes/No)",
        "• Automatic location updates when user confirms",
        "• Follow-up \"Is it lost?\" flow for missing lighters",
        "• Non-intrusive frequency (max 1-2 prompts per week)",
        "• Respects user privacy and battery life"
    ]
    
    for feature in location_features:
        story.append(Paragraph(feature, bullet_style))
    
    story.append(Paragraph(
        "This will significantly improve the lost & found success rate while keeping the experience user-friendly.",
        normal_style
    ))
    story.append(PageBreak())
    
    # Section 4
    story.append(Paragraph("4. Social Feed Architecture ✅ FUTURE-READY", section_style))
    story.append(Paragraph(
        "<b>MVP Foundation:</b> We'll structure our backend and database to support social features from day one:",
        normal_style
    ))
    story.append(Spacer(1, 0.05*inch))
    
    social_foundation = [
        "• Database Design: Activity log tables for trades, lost/found events, achievements",
        "• API Architecture: Event-driven system ready for feed generation",
        "• User Interaction Tracking: All actions logged for future social features",
        "• Scalable Infrastructure: Built to handle increased social activity"
    ]
    
    for item in social_foundation:
        story.append(Paragraph(item, bullet_style))
    
    story.append(Paragraph(
        "<b>Future Expansion Ready:</b> When you're ready to add social feeds, we'll have all the data and infrastructure in place for a smooth implementation.",
        normal_style
    ))
    story.append(Spacer(1, 0.15*inch))
    
    # Section 5
    story.append(Paragraph("5. goodmonkeys.com Landing Website ✅ ENHANCED SCOPE", section_style))
    story.append(Paragraph(
        "Your dual-purpose website concept is perfect for brand building:",
        normal_style
    ))
    story.append(Spacer(1, 0.05*inch))
    
    website_features = [
        "• Company Section: About Good Monkeys LLC, mission, team",
        "• Flick App Section: App introduction, features showcase, download buttons",
        "• Contact Form: Professional contact system with email forwarding",
        "• Responsive Design: Mobile-optimized for all devices",
        "• SEO Optimized: Ready for search engine visibility",
        "• Integrated Experience: Seamless flow between company info and app promotion"
    ]
    
    for feature in website_features:
        story.append(Paragraph(feature, bullet_style))
    
    story.append(Paragraph(
        "This creates a professional presence that serves both your company brand and Flick app marketing needs.",
        normal_style
    ))
    story.append(Spacer(1, 0.2*inch))
    
    # Updated Scope Summary
    story.append(Paragraph("📋 Updated MVP Scope Summary", section_style))
    story.append(Paragraph(
        "<b>All your requested features are now included in our $13,500 MVP development:</b>",
        normal_style
    ))
    story.append(Spacer(1, 0.1*inch))
    
    # Core Features Table
    core_features_data = [
        ['Core Features', 'Technical Foundation'],
        ['✅ iOS App with native QR scanning + web fallback', '✅ Social-ready database architecture'],
        ['✅ User profiles with privacy controls', '✅ Scalable API design'],
        ['✅ Ownership history tracking', '✅ Event logging for future features'],
        ['✅ Smart location update system', '✅ Privacy-first user controls'],
        ['✅ Lost & found workflow', ''],
        ['✅ Trading & gifting system', ''],
        ['✅ Basic gamification', ''],
        ['✅ Admin dashboard', ''],
        ['✅ Enhanced goodmonkeys.com website', '']
    ]
    
    features_table = Table(core_features_data, colWidths=[3*inch, 3*inch])
    features_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), primary_blue),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, HexColor('#f8f9fa')]),
        ('VALIGN', (0, 0), (-1, -1), 'TOP')
    ]))
    
    story.append(features_table)
    story.append(Spacer(1, 0.2*inch))
    
    # Investment Confirmation
    story.append(Paragraph("💰 Investment Confirmation", section_style))
    story.append(Paragraph(
        "<b>Total cost remains $13,500</b> - we're excited to include these valuable features at no additional cost because they align perfectly with our vision for Flick's success.",
        normal_style
    ))
    story.append(Spacer(1, 0.15*inch))
    
    # Next Steps
    story.append(Paragraph("🤝 Next Steps", section_style))
    
    next_steps = [
        "1. Confirm approval of this enhanced scope",
        "2. Schedule kickoff call to finalize technical details",
        "3. Begin development November 1-7, 2025"
    ]
    
    for step in next_steps:
        story.append(Paragraph(step, bullet_style))
    
    story.append(Spacer(1, 0.2*inch))
    
    # Closing
    story.append(Paragraph(
        "We're genuinely excited about these enhancements - they'll make Flick a truly innovative and engaging platform that users will love. Your vision for community-driven lighter tracking is going to create something special.",
        normal_style
    ))
    story.append(Spacer(1, 0.1*inch))
    story.append(Paragraph(
        "Thank you for trusting CodeFlow Studios with this exciting project!",
        normal_style
    ))
    story.append(Spacer(1, 0.15*inch))
    
    # Signature
    story.append(Paragraph("Best regards,", normal_style))
    story.append(Spacer(1, 0.1*inch))
    story.append(Paragraph("Saddam", ParagraphStyle(
        'Signature',
        parent=styles['Normal'],
        fontSize=12,
        fontName='Helvetica-Bold'
    )))
    story.append(Paragraph("Founder, CodeFlow Studios", normal_style))
    story.append(Paragraph("Pakistan", normal_style))
    story.append(Paragraph("📱 WhatsApp: +92-333-4443355", normal_style))
    story.append(Paragraph("🌐 Website: https://www.codeflowstudios.xyz/", normal_style))
    
    # Build PDF
    doc.build(story)
    print(f"✅ PDF document created: {filename}")
    return filename


if __name__ == "__main__":
    print("Creating Word document...")
    word_file = create_word_document()
    
    print("\nCreating PDF document...")
    pdf_file = create_pdf_document()
    
    print(f"\n✅ Both documents created successfully!")
    print(f"📄 Word: {word_file}")
    print(f"📄 PDF: {pdf_file}")




